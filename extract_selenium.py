import streamlit as st
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import re
import io

st.set_page_config(page_title="Web Content Scraper", page_icon="🌐", layout="wide")


def setup_driver():
    """Alternative setup - Cloud compatible"""
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--disable-blink-features=AutomationControlled")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument(
        "--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    )
    chrome_options.add_argument("--window-size=1920,1080")
    chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
    chrome_options.add_experimental_option("useAutomationExtension", False)

    # Add these lines for cloud deployment
    chrome_options.binary_location = "/usr/bin/chromium"

    try:
        # For cloud: use system chromium-driver
        service = Service("/usr/bin/chromedriver")
        driver = webdriver.Chrome(service=service, options=chrome_options)
    except:
        # Fallback for local development
        try:
            driver = webdriver.Chrome(options=chrome_options)
        except:
            from webdriver_manager.chrome import ChromeDriverManager

            service = Service(ChromeDriverManager().install())
            driver = webdriver.Chrome(service=service, options=chrome_options)

    driver.execute_script(
        "Object.defineProperty(navigator, 'webdriver', {get: () => undefined})"
    )
    return driver


def fetch_with_selenium(url, progress_callback=None):
    """EXACT SAME LOGIC AS YOUR WORKING SCRIPT"""
    driver = None
    try:
        if progress_callback:
            progress_callback("🌐 Starting Chrome browser...", 10)

        driver = setup_driver()

        if progress_callback:
            progress_callback(f"📡 Loading: {url}", 30)

        driver.get(url)

        if progress_callback:
            progress_callback(f"⏳ Waiting for page to load...", 50)

        WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.TAG_NAME, "body"))
        )

        # Additional wait for dynamic content
        time.sleep(5)

        if progress_callback:
            progress_callback("📜 Scrolling page...", 70)

        # Scroll to load lazy-loaded content
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight/2);")
        time.sleep(2)
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        time.sleep(2)
        driver.execute_script("window.scrollTo(0, 0);")
        time.sleep(1)

        html_content = driver.page_source
        page_title = driver.title

        if progress_callback:
            progress_callback("✓ Content fetched!", 90)

        return html_content, page_title

    except Exception as e:
        raise Exception(f"Error fetching content: {str(e)}")
    finally:
        if driver:
            driver.quit()


def clean_text(text):
    """Clean and normalize text"""
    if not text:
        return ""
    text = re.sub(r"[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]", "", text)
    text = re.sub(r"\s+", " ", text)
    text = text.replace("\xa0", " ").replace("\u200b", "").replace("\ufeff", "")
    return text.strip()


def extract_content(html_content):
    """Extract structured content from HTML"""
    soup = BeautifulSoup(html_content, "html.parser")

    for tag in soup(["script", "style", "noscript", "iframe", "svg", "path"]):
        tag.decompose()

    content_structure = []
    processed_texts = set()

    title = soup.find("title")
    if title:
        title_text = clean_text(title.get_text())
        if title_text and len(title_text) > 3:
            content_structure.append({"type": "title", "text": title_text})
            processed_texts.add(title_text)

    meta_desc = soup.find("meta", attrs={"name": "description"})
    if meta_desc and meta_desc.get("content"):
        desc_text = clean_text(meta_desc.get("content"))
        if desc_text and len(desc_text) > 10:
            content_structure.append({"type": "paragraph", "text": desc_text})
            processed_texts.add(desc_text)

    main_content = (
        soup.find("main")
        or soup.find("article")
        or soup.find(id=re.compile(r"content|main", re.I))
        or soup.find(class_=re.compile(r"content|main|wrapper", re.I))
        or soup.find("body")
    )

    for element in main_content.find_all(
        [
            "h1",
            "h2",
            "h3",
            "h4",
            "h5",
            "h6",
            "p",
            "li",
            "td",
            "blockquote",
            "div",
            "span",
            "a",
        ]
    ):
        text = clean_text(element.get_text(strip=True))

        if not text or len(text) < 10 or text in processed_texts:
            continue

        if len(text) < 20 and element.name in ["a", "span", "div"]:
            continue

        processed_texts.add(text)

        if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            content_structure.append({"type": element.name, "text": text})
        elif len(text) > 30:
            content_structure.append({"type": "paragraph", "text": text})

    return content_structure


def create_word_document(content_structure):
    """Create a formatted Word document"""
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for item in content_structure:
        try:
            text = item["text"]

            if item["type"] == "title":
                heading = doc.add_heading(text, level=0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
            elif item["type"] == "h1":
                doc.add_heading(text, level=1)
            elif item["type"] == "h2":
                doc.add_heading(text, level=2)
            elif item["type"] == "h3":
                doc.add_heading(text, level=3)
            elif item["type"] == "h4":
                doc.add_heading(text, level=4)
            elif item["type"] == "paragraph":
                doc.add_paragraph(text)
        except:
            continue

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io


def create_text_file(content_structure):
    """Create a plain text file"""
    text_content = []

    for item in content_structure:
        text = item["text"]

        if item["type"] == "title":
            text_content.append("=" * 80)
            text_content.append(text.upper())
            text_content.append("=" * 80)
            text_content.append("")
        elif item["type"] in ["h1", "h2", "h3", "h4"]:
            level = int(item["type"][1])
            text_content.append("")
            text_content.append("#" * level + " " + text)
            text_content.append("")
        elif item["type"] == "paragraph":
            text_content.append(text)
            text_content.append("")

    return "\n".join(text_content)


# ==================== STREAMLIT UI ====================

st.title("🌐 Web Content Scraper")
st.markdown("Extract all text content from any website with proper formatting")

st.divider()

url = st.text_input(
    "Enter Website URL:",
    placeholder="https://example.com",
    help="Enter the complete URL including https://",
)

col1, col2 = st.columns([1, 4])
with col1:
    scrape_button = st.button(
        "🚀 Scrape Website", type="primary", use_container_width=True
    )

st.divider()

if scrape_button:
    if not url:
        st.error("❌ Please enter a URL")
    elif not url.startswith(("http://", "https://")):
        st.error("❌ URL must start with http:// or https://")
    else:
        progress_bar = st.progress(0)
        status_text = st.empty()

        def update_progress(message, percent):
            progress_bar.progress(percent)
            status_text.text(message)

        try:
            # Fetch content
            html_content, page_title = fetch_with_selenium(url, update_progress)

            update_progress("Extracting content...", 95)
            content_structure = extract_content(html_content)

            update_progress("Complete!", 100)
            time.sleep(0.5)
            progress_bar.empty()
            status_text.empty()

            if not content_structure:
                st.error("❌ No content extracted from the website")
            else:
                st.success(
                    f"✅ Successfully extracted {len(content_structure)} content blocks!"
                )

                # Display preview
                st.subheader("📄 Content Preview")
                with st.expander("Click to view extracted content", expanded=True):
                    for i, item in enumerate(content_structure[:10], 1):
                        if item["type"] == "title":
                            st.markdown(f"### {item['text']}")
                        elif item["type"] in ["h1", "h2", "h3"]:
                            st.markdown(f"**{item['text']}**")
                        else:
                            preview = item["text"][:200]
                            st.write(f"{preview}...")

                    if len(content_structure) > 10:
                        st.info(f"... and {len(content_structure) - 10} more items")

                st.divider()

                # Download section
                st.subheader("💾 Download Files")

                col1, col2 = st.columns(2)

                with col1:
                    doc_io = create_word_document(content_structure)
                    st.download_button(
                        label="📄 Download Word Document (.docx)",
                        data=doc_io,
                        file_name="scraped_content.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )

                with col2:
                    text_content = create_text_file(content_structure)
                    st.download_button(
                        label="📝 Download Text File (.txt)",
                        data=text_content,
                        file_name="scraped_content.txt",
                        mime="text/plain",
                        use_container_width=True,
                    )

                st.info(
                    "💡 **Tip:** Upload the .docx file to Google Drive and open with Google Docs!"
                )

        except Exception as e:
            progress_bar.empty()
            status_text.empty()
            st.error(f"❌ Error: {str(e)}")

st.divider()
st.markdown(
    """
<div style='text-align: center; color: gray; padding: 20px;'>
    <p>Web Content Scraper v1.0</p>
</div>
""",
    unsafe_allow_html=True,
)
